import os
from fastapi import FastAPI, Request, UploadFile, File, HTTPException
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
from summarizer import summarize_hindi

app = FastAPI()

# Config
UPLOAD_FOLDER = "uploads"
SUMMARY_FOLDER = "summaries"
SECRET_KEY = "hindi_summarizer_secret_key"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(SUMMARY_FOLDER, exist_ok=True)

# Serve static & templates (same as Flask render_template)
templates = Jinja2Templates(directory="templates")
app.mount("/static", StaticFiles(directory="static"), name="static")


@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "flash": None})


@app.post("/", response_class=HTMLResponse)
async def upload_file(request: Request, doc_file: UploadFile = File(...)):
    # Check file
    if not doc_file.filename.endswith(".docx"):
        return templates.TemplateResponse(
            "index.html",
            {"request": request, "flash": "Please upload a valid .docx file"}
        )

    try:
        # Save uploaded file
        filepath = os.path.join(UPLOAD_FOLDER, doc_file.filename)
        with open(filepath, "wb") as f:
            f.write(await doc_file.read())

        # Read DOCX
        doc = Document(filepath)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text])

        if not full_text.strip():
            return templates.TemplateResponse(
                "index.html",
                {"request": request, "flash": "The uploaded document appears to be empty"}
            )

        # Summarize
        summary_text = summarize_hindi(full_text)

        # Write summary DOCX
        summary_doc = Document()
        heading = summary_doc.add_heading("", level=1)
        heading_run = heading.add_run("सारांश")
        heading_run.font.name = "Nirmala UI"

        for line in summary_text.split("\n"):
            if line.strip():
                para = summary_doc.add_paragraph("")
                run = para.add_run(line)
                run.font.name = "Nirmala UI"

        summary_filename = f"summary_{doc_file.filename}"
        summary_path = os.path.join(SUMMARY_FOLDER, summary_filename)
        summary_doc.save(summary_path)

        return RedirectResponse(url=f"/download/{summary_filename}", status_code=303)

    except Exception as e:
        return templates.TemplateResponse(
            "index.html",
            {"request": request, "flash": f"Error processing file: {str(e)}"}
        )


@app.get("/download/{filename}")
async def download_summary(filename: str):
    path = os.path.join(SUMMARY_FOLDER, filename)
    if os.path.exists(path):
        return FileResponse(
            path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename
        )
    raise HTTPException(status_code=404, detail="Summary file not found")


@app.exception_handler(404)
async def custom_404_handler(request: Request, exc):
    return templates.TemplateResponse("404.html", {"request": request}, status_code=404)


@app.exception_handler(500)
async def custom_500_handler(request: Request, exc):
    return templates.TemplateResponse("500.html", {"request": request}, status_code=500)


# For local run
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)
